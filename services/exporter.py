import os
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor


class ResumeExporter:
    """Export merged resume to PDF or DOCX"""
    
    def __init__(self):
        self.output_dir = "exports"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export(self, resume_data: Dict[str, Any], format: str) -> str:
        """Export resume in specified format"""
        if format == "pdf":
            return self._export_pdf(resume_data)
        elif format == "docx":
            return self._export_docx(resume_data)
        else:
            raise ValueError("Unsupported format")
    
    def _export_pdf(self, resume_data: Dict[str, Any]) -> str:
        """Export to PDF using ReportLab"""
        output_path = os.path.join(self.output_dir, "merged_resume.pdf")
        
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            spaceBefore=12
        )
        
        # Personal Info
        personal = resume_data.get('personal_info', {})
        if personal.get('name'):
            story.append(Paragraph(personal['name'], title_style))
        
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), styles['Normal']))
        
        story.append(Spacer(1, 0.2*inch))
        
        # Skills
        skills = resume_data.get('skills', [])
        if skills:
            story.append(Paragraph("SKILLS", heading_style))
            skill_names = [s['name'] for s in skills[:20]]  # Top 20 skills
            story.append(Paragraph(", ".join(skill_names), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Experience
        experiences = resume_data.get('experience', [])
        if experiences:
            story.append(Paragraph("EXPERIENCE", heading_style))
            for exp in experiences:
                title_company = f"<b>{exp['title']}</b> - {exp['company']}"
                story.append(Paragraph(title_company, styles['Normal']))
                
                for desc in exp.get('description', [])[:3]:  # Limit descriptions
                    story.append(Paragraph(f"• {desc}", styles['Normal']))
                
                story.append(Spacer(1, 0.1*inch))
        
        # Education
        education = resume_data.get('education', [])
        if education:
            story.append(Paragraph("EDUCATION", heading_style))
            for edu in education:
                edu_text = f"<b>{edu['degree']}</b> - {edu['institution']}"
                story.append(Paragraph(edu_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def _export_docx(self, resume_data: Dict[str, Any]) -> str:
        """Export to DOCX"""
        output_path = os.path.join(self.output_dir, "merged_resume.docx")
        
        doc = Document()
        
        # Personal Info
        personal = resume_data.get('personal_info', {})
        if personal.get('name'):
            name_para = doc.add_heading(personal['name'], level=1)
            name_para.runs[0].font.size = Pt(24)
        
        contact_info = []
        if personal.get('email'):
            contact_info.append(personal['email'])
        if personal.get('phone'):
            contact_info.append(personal['phone'])
        if contact_info:
            doc.add_paragraph(" | ".join(contact_info))
        
        # Skills
        skills = resume_data.get('skills', [])
        if skills:
            doc.add_heading('SKILLS', level=2)
            skill_names = [s['name'] for s in skills[:20]]
            doc.add_paragraph(", ".join(skill_names))
        
        # Experience
        experiences = resume_data.get('experience', [])
        if experiences:
            doc.add_heading('EXPERIENCE', level=2)
            for exp in experiences:
                doc.add_paragraph(f"{exp['title']} - {exp['company']}", style='Heading 3')
                for desc in exp.get('description', [])[:3]:
                    doc.add_paragraph(f"• {desc}")
        
        # Education
        education = resume_data.get('education', [])
        if education:
            doc.add_heading('EDUCATION', level=2)
            for edu in education:
                doc.add_paragraph(f"{edu['degree']} - {edu['institution']}")
        
        doc.save(output_path)
        return output_path
